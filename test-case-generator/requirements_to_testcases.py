#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
需求文档转测试用例生成器
将doc/docx格式的需求文档转换为xlsx格式的测试用例
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Set
import re

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    print("需要安装 openpyxl: pip install openpyxl")
    sys.exit(1)


class RequirementParser:
    """需求文档解析器"""
    
    def __init__(self, doc_path: str):
        """
        初始化解析器
        
        Args:
            doc_path: doc/docx文件路径
        """
        self.doc_path = doc_path
        self.requirements = []
        self.document = None
        
    def parse(self) -> List[Dict[str, str]]:
        """
        解析需求文档
        
        Returns:
            需求列表，每个需求包含id、标题、描述等信息
        """
        try:
            self.document = Document(self.doc_path)
        except Exception as e:
            print(f"错误：无法打开文档 {self.doc_path}")
            print(f"详情：{e}")
            return []
        
        # 提取文本内容
        full_text = self._extract_text()
        
        # 解析需求
        self.requirements = self._parse_requirements(full_text)
        
        return self.requirements
    
    def _extract_text(self) -> str:
        """
        从document对象中提取所有文本
        
        Returns:
            合并后的文本内容
        """
        text_parts = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 处理表格中的文本
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def _parse_requirements(self, text: str) -> List[Dict[str, str]]:
        """
        解析需求文本
        
        Args:
            text: 文档文本内容
            
        Returns:
            需求列表
        """
        requirements = []

        # 按行分割
        lines = text.split('\n')

        # 锚定到行首，避免“包含关键词”误判为新需求
        title_pattern = re.compile(
            r'^(?:需求\s*[\dA-Za-z_-]*|功能\s*[\dA-Za-z_-]*|FR\s*[-_\dA-Za-z]*|Requirement\s*[-_\dA-Za-z]*)\s*[:：]?.*$',
            re.IGNORECASE
        )
        desc_pattern = re.compile(r'^(描述|说明|Description)\s*[:：]\s*(.*)$', re.IGNORECASE)
        precondition_pattern = re.compile(r'^(前置条件|Precondition|前提)\s*[:：]\s*(.*)$', re.IGNORECASE)
        step_pattern = re.compile(r'^(步骤|操作|Step|Action)\s*[:：]\s*(.*)$', re.IGNORECASE)
        expected_pattern = re.compile(r'^(预期结果|预期|Expected\s*Result|Expected|Result)\s*[:：]\s*(.*)$', re.IGNORECASE)

        current_requirement = None
        last_field = None

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue

            if title_pattern.match(line):
                if current_requirement:
                    requirements.append(current_requirement)

                current_requirement = {
                    'title': line,
                    'description': '',
                    'precondition': '系统正常运行',
                    'steps': [],
                    'expected_result': ''
                }
                last_field = 'title'
                continue

            if not current_requirement:
                continue

            desc_match = desc_pattern.match(line)
            if desc_match:
                current_requirement['description'] = desc_match.group(2).strip()
                last_field = 'description'
                continue

            precondition_match = precondition_pattern.match(line)
            if precondition_match:
                value = precondition_match.group(2).strip()
                current_requirement['precondition'] = value or '系统正常运行'
                last_field = 'precondition'
                continue

            step_match = step_pattern.match(line)
            if step_match:
                value = step_match.group(2).strip()
                if value:
                    current_requirement['steps'].append(value)
                last_field = 'steps'
                continue

            expected_match = expected_pattern.match(line)
            if expected_match:
                current_requirement['expected_result'] = expected_match.group(2).strip()
                last_field = 'expected_result'
                continue

            # 处理多行字段续写
            if last_field == 'steps' and current_requirement['steps']:
                current_requirement['steps'][-1] += f'\n{line}'
            elif last_field == 'expected_result' and current_requirement['expected_result']:
                current_requirement['expected_result'] += f'\n{line}'
            else:
                if current_requirement['description']:
                    current_requirement['description'] += f'\n{line}'
                else:
                    current_requirement['description'] = line
                last_field = 'description'

        # 添加最后一个需求
        if current_requirement:
            requirements.append(current_requirement)

        return requirements


class TestCaseGenerator:
    """测试用例生成器"""
    
    def __init__(self, requirements: List[Dict[str, str]]):
        """
        初始化生成器
        
        Args:
            requirements: 需求列表
        """
        self.requirements = requirements
        self.test_cases = []
        
    def generate(self) -> List[Dict[str, str]]:
        """
        根据需求生成测试用例
        
        Returns:
            测试用例列表
        """
        for idx, req in enumerate(self.requirements, 1):
            # 基于覆盖策略为每个需求生成多类型测试用例
            test_cases = self._generate_cases_for_requirement(idx, req)
            self.test_cases.extend(test_cases)
        
        return self.test_cases
    
    def _generate_cases_for_requirement(self, req_id: int, req: Dict[str, str]) -> List[Dict[str, str]]:
        """为单个需求生成覆盖驱动的测试用例。"""
        title = req.get('title', '').strip() or f'需求{req_id}'
        description = req.get('description', '').strip() or title
        precondition = req.get('precondition', '').strip() or '系统正常运行'
        expected_result = req.get('expected_result', '').strip() or '功能执行成功'
        step_items = self._extract_step_items(req.get('steps', []))
        steps_text = '\n'.join(step_items) if step_items else '执行对应功能'
        keywords = self._collect_keywords(title, description, precondition, steps_text, expected_result)

        cases = []
        seq = 1
        seen_titles: Set[str] = set()

        def add_case(case_title: str, case_desc: str, case_steps: str, case_expected: str,
                     priority: str, test_type: str):
            nonlocal seq
            if case_title in seen_titles:
                return
            seen_titles.add(case_title)
            cases.append({
                'test_case_id': f'TC_{req_id}_{seq:03d}',
                'requirement_id': f'REQ_{req_id}',
                'title': case_title,
                'description': case_desc,
                'precondition': precondition,
                'steps': case_steps,
                'expected_result': case_expected,
                'priority': priority,
                'test_type': test_type
            })
            seq += 1

        add_case(f'{title} - 正向主流程', description, steps_text, expected_result, '高', '功能测试')

        for idx, step in enumerate(step_items, 1):
            add_case(
                f'{title} - 步骤覆盖{idx}',
                f'验证关键步骤“{step}”执行正确且不破坏主流程',
                f'1. 执行步骤：{step}\n2. 使用合法输入继续流程\n3. 校验步骤结果',
                f'步骤“{step}”执行成功且后续流程可继续',
                '高',
                '流程覆盖测试'
            )

        generic_scenarios = [
            ('边界值-最小值', '在最小边界条件下执行流程', '系统可在最小边界值下正确处理', '中', '边界值测试'),
            ('边界值-最大值', '在最大边界条件下执行流程', '系统可在最大边界值下正确处理', '中', '边界值测试'),
            ('边界值-临界点', '在临界值附近执行流程', '系统对临界值判断准确，无越界错误', '中', '边界值测试'),
            ('异常-必填项缺失', '缺失关键输入后执行流程', '系统阻止提交并提示必填项', '高', '异常处理测试'),
            ('异常-非法格式输入', '输入非法格式数据后执行流程', '系统正确拦截并提示格式错误', '高', '异常处理测试'),
            ('异常-超范围输入', '输入超出允许范围的数据后执行流程', '系统拦截超范围输入并提示原因', '中', '异常处理测试'),
            ('异常-重复提交', '短时间内重复触发提交动作', '系统保持幂等，避免重复数据写入', '高', '特殊场景测试'),
            ('异常-服务超时', '模拟接口响应超时后继续操作', '系统提示超时并允许重试', '中', '异常处理测试'),
            ('特殊-会话超时恢复', '会话过期后重新发起流程', '系统引导重新认证并可恢复操作', '中', '特殊场景测试')
        ]

        for scenario_name, scenario_desc, scenario_expected, priority, case_type in generic_scenarios:
            add_case(
                f'{title} - {scenario_name}',
                f'{description}（{scenario_desc}）',
                f'1. 准备场景：{scenario_desc}\n2. 执行需求流程\n3. 观察系统响应',
                scenario_expected,
                priority,
                case_type
            )

        if keywords & {'登录', '认证', '密码', '验证码', 'token'}:
            add_case(
                f'{title} - 安全专项-凭证错误',
                f'{description}（安全专项）',
                '1. 使用错误凭证执行流程\n2. 连续尝试\n3. 观察风控与提示',
                '系统拒绝非法访问并提供安全提示，不泄露敏感信息',
                '高',
                '安全性测试'
            )

        if keywords & {'权限', '角色', '管理员', '授权'}:
            add_case(
                f'{title} - 权限专项-越权访问',
                f'{description}（权限专项）',
                '1. 使用低权限账号访问高权限功能\n2. 尝试执行关键操作',
                '系统拦截越权操作并记录审计日志',
                '高',
                '权限测试'
            )

        if keywords & {'并发', '同时', '抢购', '库存', '高并发', '下单', '支付', '提交', '保存'}:
            add_case(
                f'{title} - 并发专项-竞争条件',
                f'{description}（并发专项）',
                '1. 多用户同时触发同一操作\n2. 观察结果与数据一致性',
                '系统正确处理并发冲突，保证数据一致性',
                '高',
                '并发测试'
            )

        if keywords & {'网络', '超时', '断网', '重试'}:
            add_case(
                f'{title} - 稳定性专项-网络抖动恢复',
                f'{description}（稳定性专项）',
                '1. 执行流程中模拟网络中断\n2. 恢复网络后继续操作\n3. 校验状态恢复',
                '系统可恢复到可用状态，避免脏数据与重复提交',
                '中',
                '稳定性测试'
            )

        if keywords & {'导出', '上传', '下载', '文件', '附件'}:
            add_case(
                f'{title} - 文件专项-格式与大小校验',
                f'{description}（文件处理专项）',
                '1. 分别上传/导出合法与非法文件\n2. 校验格式、大小、编码等边界',
                '系统正确处理合法文件并拒绝不符合约束的文件',
                '中',
                '文件处理测试'
            )

        if keywords & {'ios', 'android', 'app', '浏览器', '小程序', '客户端'}:
            add_case(
                f'{title} - 兼容性专项-多端一致性',
                f'{description}（兼容性专项）',
                '1. 在不同终端执行同一流程\n2. 比对页面和结果一致性',
                '核心功能在目标终端表现一致，差异在可接受范围内',
                '中',
                '兼容性测试'
            )

        if keywords & {'查询', '搜索', '筛选', '列表', '分页'}:
            add_case(
                f'{title} - 查询专项-空结果与性能',
                f'{description}（查询专项）',
                '1. 查询空结果\n2. 查询大数据量并分页\n3. 校验返回性能与准确性',
                '空结果有清晰提示，分页准确且性能达标',
                '中',
                '特殊场景测试'
            )

        return cases

    def _extract_step_items(self, steps: List[str]) -> List[str]:
        """将步骤文本拆分为关键步骤列表。"""
        if not steps:
            return []

        items = []
        for raw in steps:
            text = raw.strip()
            if not text:
                continue

            parts = re.split(r'\s*(?=\d+[\.|、\)])', text)
            for part in parts:
                cleaned = re.sub(r'^\d+[\.|、\)]\s*', '', part).strip()
                if cleaned:
                    items.append(cleaned)

        if not items:
            items = [step.strip() for step in steps if step.strip()]

        seen = set()
        deduplicated = []
        for item in items:
            if item not in seen:
                seen.add(item)
                deduplicated.append(item)
        return deduplicated

    def _collect_keywords(self, *text_parts: str) -> Set[str]:
        """提取需求关键词用于扩展专项测试。"""
        combined = ' '.join(part for part in text_parts if part)
        lower = combined.lower()
        candidate_keywords = {
            '登录', '认证', '密码', '验证码', 'token',
            '权限', '角色', '管理员', '授权',
            '并发', '同时', '抢购', '库存', '高并发', '下单', '支付', '提交', '保存',
            '网络', '超时', '断网', '重试',
            '导出', '上传', '下载', '文件', '附件',
            'ios', 'android', 'app', '浏览器', '小程序', '客户端',
            '查询', '搜索', '筛选', '列表', '分页'
        }

        hits = set()
        for keyword in candidate_keywords:
            if keyword in combined or keyword in lower:
                hits.add(keyword)
        return hits


class ExcelExporter:
    """Excel导出器"""

    # 定义列头和宽度
    COLUMNS = {
        'test_case_id': ('测试用例ID', 12),
        'requirement_id': ('需求ID', 12),
        'title': ('测试用例标题', 25),
        'description': ('用例描述', 20),
        'precondition': ('前置条件', 20),
        'steps': ('测试步骤', 30),
        'expected_result': ('预期结果', 30),
        'priority': ('优先级', 10),
        'test_type': ('测试类型', 12)
    }

    def __init__(self, test_cases: List[Dict[str, str]], output_path: str):
        """
        初始化导出器

        Args:
            test_cases: 测试用例列表
            output_path: 输出文件路径
        """
        self.test_cases = test_cases
        self.output_path = output_path
        self.workbook = None
        self.worksheet = None
        
    def export(self) -> bool:
        """
        导出测试用例到Excel
        
        Returns:
            是否导出成功
        """
        try:
            self.workbook = openpyxl.Workbook()
            self.worksheet = self.workbook.active
            self.worksheet.title = '测试用例'
            
            # 写入表头
            self._write_headers()
            
            # 写入数据
            self._write_data()
            
            # 调整列宽和行高
            self._adjust_format()
            
            # 保存文件
            self.workbook.save(self.output_path)
            print(f"✓ 测试用例已成功导出到: {self.output_path}")
            return True
            
        except Exception as e:
            print(f"✗ 导出Excel失败: {e}")
            return False
    
    def _write_headers(self):
        """写入表头"""
        # 定义样式
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_font = Font(bold=True, color='FFFFFF', size=11)
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for col_idx, (key, (label, width)) in enumerate(self.COLUMNS.items(), 1):
            cell = self.worksheet.cell(row=1, column=col_idx)
            cell.value = label
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment
            cell.border = border
            self.worksheet.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = width
        
        # 冻结表头
        self.worksheet.freeze_panes = 'A2'
    
    def _write_data(self):
        """写入测试用例数据"""
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row_idx, test_case in enumerate(self.test_cases, 2):
            for col_idx, key in enumerate(self.COLUMNS.keys(), 1):
                cell = self.worksheet.cell(row=row_idx, column=col_idx)
                cell.value = test_case.get(key, '')
                cell.border = border
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
                
                # 设置行高
                if test_case.get(key, ''):
                    line_count = str(test_case.get(key, '')).count('\n') + 1
                    self.worksheet.row_dimensions[row_idx].height = max(20, line_count * 15)
    
    def _adjust_format(self):
        """调整格式"""
        # 为所有有内容的行设置对齐方式
        for row in self.worksheet.iter_rows():
            for cell in row:
                if cell.value:
                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)


class TestCaseGeneratorApp:
    """测试用例生成应用主类"""
    
    def __init__(self):
        self.parser = None
        self.generator = None
        self.exporter = None
        
    def run(self, input_doc: str, output_excel: str = None) -> bool:
        """
        执行转换流程
        
        Args:
            input_doc: 输入的doc/docx文件路径
            output_excel: 输出的xlsx文件路径（如果为None则自动生成）
            
        Returns:
            是否执行成功
        """
        # 验证输入文件
        if not os.path.exists(input_doc):
            print(f"✗ 错误：文件不存在 - {input_doc}")
            return False
        
        if not input_doc.lower().endswith('.docx'):
            print(f"✗ 错误：不支持的文件格式，请使用docx文件")
            return False
        
        # 确定输出文件路径
        if output_excel is None:
            base_name = Path(input_doc).stem
            output_excel = f"{base_name}_测试用例.xlsx"
        
        print(f"\n{'='*60}")
        print(f"需求文档转测试用例生成器")
        print(f"{'='*60}")
        print(f"输入文件: {input_doc}")
        print(f"输出文件: {output_excel}")
        print(f"{'='*60}\n")
        
        # 1. 解析需求文档
        print("1️⃣  解析需求文档...")
        self.parser = RequirementParser(input_doc)
        requirements = self.parser.parse()
        
        if not requirements:
            print("✗ 未能识别到需求内容")
            return False
        
        print(f"✓ 识别到 {len(requirements)} 个需求")
        for i, req in enumerate(requirements, 1):
            print(f"   {i}. {req['title'][:50]}")
        
        # 2. 生成测试用例
        print("\n2️⃣  生成测试用例...")
        self.generator = TestCaseGenerator(requirements)
        test_cases = self.generator.generate()
        
        print(f"✓ 生成了 {len(test_cases)} 个测试用例")
        
        # 3. 导出到Excel
        print("\n3️⃣  导出到Excel...")
        self.exporter = ExcelExporter(test_cases, output_excel)
        success = self.exporter.export()
        
        if success:
            print(f"\n{'='*60}")
            print("✅ 转换完成！")
            print(f"{'='*60}\n")
        
        return success


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='将需求文档(docx)转换为测试用例(xlsx)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python requirements_to_testcases.py requirements.docx
  python requirements_to_testcases.py requirements.docx -o test_cases.xlsx
        """
    )
    
    parser.add_argument('input', help='输入的需求文档(docx格式)')
    parser.add_argument('-o', '--output', help='输出的Excel文件路径(可选，默认自动生成)', 
                       default=None)
    
    args = parser.parse_args()
    
    # 运行应用
    app = TestCaseGeneratorApp()
    success = app.run(args.input, args.output)
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
